"""
One-off builder: docs/TIER3-COURSE-SUBMISSION.md -> docs/TIER3-COURSE-SUBMISSION.docx
Requires: pip install python-docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Install dependency: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "TIER3-COURSE-SUBMISSION.md"
OUT_PATH = ROOT / "docs" / "TIER3-COURSE-SUBMISSION.docx"


def add_para_with_bold(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    for chunk in re.split(r"(\*\*.+?\*\*)", text):
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) >= 4:
            r = p.add_run(chunk[2:-2])
            r.bold = True
        elif chunk:
            p.add_run(chunk)


def is_table_sep(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [c.strip() for c in line.split("|") if c.strip() != ""]
    if len(cells) < 2:
        return False
    for c in cells:
        if not re.match(r"^:?-{2,}:?$", c):
            return False
    return True


def parse_table_row(line: str) -> list[str]:
    parts = line.strip().split("|")
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [c.strip() for c in parts]


def strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def add_markdown_table(doc: Document, header: list[str], body: list[list[str]]) -> None:
    if not header:
        return
    num_cols = max(len(header), max((len(r) for r in body), default=0))
    table = doc.add_table(rows=1 + len(body), cols=num_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header):
        if j < num_cols:
            hdr_cells[j].text = strip_md_bold(h)
    for i, row in enumerate(body):
        for j in range(num_cols):
            val = row[j] if j < len(row) else ""
            table.rows[i + 1].cells[j].text = strip_md_bold(val)


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Inches(11)
    sect.page_width = Inches(8.5)

    i = 0
    table_buf: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        sep_idx: int | None = None
        for idx, L in enumerate(table_buf):
            if is_table_sep(L):
                sep_idx = idx
                break
        if sep_idx is None or sep_idx == 0:
            table_buf = []
            return
        header = parse_table_row(table_buf[0])
        body = [parse_table_row(L) for L in table_buf[sep_idx + 1 :]]
        add_markdown_table(doc, header, body)
        table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            flush_table()
            i += 1
            continue

        if line.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if line.startswith("# ") and not line.startswith("##"):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip(), style="Intense Quote")
            i += 1
            continue

        if re.match(r"^[\-*] \[ \]", line.strip()):
            doc.add_paragraph(line.strip()[4:].strip(), style="List Paragraph")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            doc.add_paragraph(line.strip(), style="List Number")
            i += 1
            continue

        if re.match(r"^[\-*]\s+", line.strip()):
            doc.add_paragraph(line.strip()[2:].strip(), style="List Bullet")
            i += 1
            continue

        if line.strip():
            add_para_with_bold(doc, line)
        i += 1

    flush_table()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
